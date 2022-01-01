import pickle
import docx
  
doc = docx.Document()
  
doc.add_heading('Table 2', 0)

table = doc.add_table(rows=9 , cols=5)

row = table.rows[0].cells
row[0].text = 'Algo Name'
row[1].text = 'Accuracy'
row[2].text = 'Specificity'
row[3].text = 'Sensitivity'
row[4].text = 'AUC'


with open('test_results.pkl', 'rb') as f:
    while True:
        try:
            a = pickle.load(f)
            keys_ = a.__dict__['folds'][0].__dict__ 
            keys_["auc"] = a.auc_avg()
            row = table.add_row().cells
            row[0].text = keys_["trained_classifier"].__class__.__name__
            row[1].text = str("{:.2f}".format(keys_["accuracy"]))
            row[2].text = str("{:.2f}".format(keys_["specificity"]))
            row[3].text = str("{:.2f}".format(keys_["sensitivity"]))
            row[4].text = str("{:.2f}".format(keys_["auc"]))
        except EOFError:
                break 
table.style = 'Medium Grid 3 Accent 1'
doc.save("table2.docx")