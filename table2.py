import pickle
import docx
  
doc = docx.Document()
doc.add_heading('Table 2', 0)
table = doc.add_table(rows=1 , cols=6)

row = table.rows[0].cells
row[0].text = 'Algo Name'
row[1].text = "Task"
row[2].text = 'Accuracy'
row[3].text = 'Specificity'
row[4].text = 'Sensitivity'
row[5].text = 'AUC'

with open('test_results.pkl', 'rb') as f:
    while True:
        try:
            a = pickle.load(f)
            tar_ = a.__dict__["prediction_target"]
            keys_ = a.__dict__['folds'][0].__dict__ 
            row = table.add_row().cells
            row[0].text = a.__dict__['folds'][0].__dict__["trained_classifier"].__class__.__name__
            row[1].text = tar_
            row[2].text = str("{:.2f}".format(a.acc_avg())) # accuracy
            row[3].text = str("{:.2f}".format(a.specificity_avg())) # specificity
            row[4].text = str("{:.2f}".format(a.sensitivity_avg())) # sensitivity
            row[5].text = str("{:.2f}".format(a.auc_avg())) # auc
        except EOFError:
                break 
table.style = 'Medium Grid 3 Accent 1'
doc.save("table2.docx")